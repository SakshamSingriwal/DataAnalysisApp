import pandas as pd
import numpy as np
import streamlit as st
import io
import os
import zipfile
import sqlite3
import traceback
import json
from datetime import datetime

# ════════════════════════════════════════════════════════════════
# PAGE CONFIG — must be the very first Streamlit command
# ════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="DataIQ Pro — Enterprise Analytics",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

/* ── Reset & base ── */
*, *::before, *::after { box-sizing: border-box; }
html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
#MainMenu, footer, header { visibility: hidden; }
.main { background: #060810; }
.block-container { padding: 0 2rem 4rem 2rem !important; max-width: 100% !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: #060810; }
::-webkit-scrollbar-thumb { background: #1e293b; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #6366f1; }

/* ── Metrics ── */
[data-testid="metric-container"] {
    background: linear-gradient(145deg, #0f172a, #111827);
    border: 1px solid #1e293b;
    border-radius: 14px;
    padding: 16px 20px;
    transition: all .25s ease;
    position: relative;
    overflow: hidden;
}
[data-testid="metric-container"]::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 2px;
    background: linear-gradient(90deg, #6366f1, #8b5cf6, #ec4899);
}
[data-testid="metric-container"]:hover {
    border-color: #6366f1;
    transform: translateY(-2px);
    box-shadow: 0 8px 32px rgba(99,102,241,.2);
}
[data-testid="metric-label"] {
    color: #64748b !important; font-size: .72em !important;
    font-weight: 600 !important; letter-spacing: .8px !important;
    text-transform: uppercase !important;
}
[data-testid="metric-value"] {
    color: #f1f5f9 !important; font-size: 1.6em !important;
    font-weight: 800 !important;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    background: #0f172a;
    border-radius: 14px;
    padding: 6px;
    gap: 4px;
    border: 1px solid #1e293b;
    flex-wrap: wrap;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 10px;
    color: #64748b;
    padding: 9px 22px;
    font-weight: 600;
    font-size: .82em;
    letter-spacing: .3px;
    transition: all .2s;
    white-space: nowrap;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
    color: #fff !important;
    box-shadow: 0 4px 16px rgba(99,102,241,.35) !important;
}

/* ── Expander — fix text overlap ── */
.stExpander {
    background: #0f172a !important;
    border: 1px solid #1e293b !important;
    border-radius: 14px !important;
    margin: 6px 0 !important;
    overflow: hidden !important;
}
.stExpander:hover { border-color: #6366f1 !important; }
.stExpander details summary {
    color: #e2e8f0 !important;
    font-weight: 600 !important;
    font-size: .88em !important;
    padding: 14px 18px !important;
    display: flex !important;
    align-items: center !important;
    gap: 8px !important;
    cursor: pointer !important;
    list-style: none !important;
    background: transparent !important;
}
/* Ensure summary text never wraps on top of the arrow */
.stExpander details summary::marker,
.stExpander details summary::-webkit-details-marker { display: none; }

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #6366f1, #8b5cf6) !important;
    color: #fff !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 11px 28px !important;
    font-weight: 700 !important;
    font-size: .85em !important;
    letter-spacing: .4px !important;
    width: 100% !important;
    transition: all .2s !important;
    box-shadow: 0 4px 16px rgba(99,102,241,.3) !important;
}
.stButton > button:hover {
    opacity: .9 !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 8px 24px rgba(99,102,241,.45) !important;
}

/* ── Select / multiselect / inputs ── */
.stSelectbox > div > div,
.stMultiSelect > div > div,
.stTextInput > div > div > input,
.stNumberInput > div > div > input {
    background: #0f172a !important;
    border: 1px solid #1e293b !important;
    border-radius: 10px !important;
    color: #e2e8f0 !important;
}
.stSelectbox > div > div:focus-within,
.stTextInput > div > div > input:focus {
    border-color: #6366f1 !important;
    box-shadow: 0 0 0 3px rgba(99,102,241,.15) !important;
}
/* Selectbox option text colour */
[data-baseweb="select"] [data-testid="stSelectboxValue"],
[data-baseweb="select"] span { color: #e2e8f0 !important; }

/* ── DataFrames ── */
div[data-testid="stDataFrame"] {
    border-radius: 12px !important;
    overflow: hidden !important;
    border: 1px solid #1e293b !important;
}

/* ── File uploader — NO overlap fix ── */
[data-testid="stFileUploader"] {
    background: #0f172a !important;
    border: 2px dashed #1e293b !important;
    border-radius: 16px !important;
    padding: 8px 16px !important;
    transition: all .25s !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: #6366f1 !important;
    background: #0f1629 !important;
}
/* Prevent label from overlapping the drop zone */
[data-testid="stFileUploaderDropzoneInstructions"] {
    pointer-events: none;
}

/* ── Slider ── */
.stSlider > div > div > div { background: #6366f1 !important; }

/* ── Code blocks ── */
.stCode, code {
    background: #0f172a !important;
    border: 1px solid #1e293b !important;
    border-radius: 8px !important;
}

/* ── Label colours ── */
label, .stSelectbox label, .stMultiSelect label,
.stTextInput label, .stSlider label {
    color: #64748b !important;
    font-size: .78em !important;
    font-weight: 600 !important;
    letter-spacing: .4px !important;
}

/* ════ Custom component classes ════ */
.iq-page-header {
    padding: 36px 0 16px;
    border-bottom: 1px solid #1e293b;
    margin-bottom: 28px;
}
.iq-logo {
    font-size: 2.8em;
    font-weight: 900;
    background: linear-gradient(135deg, #6366f1, #8b5cf6, #ec4899, #f59e0b);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: -1.5px;
    line-height: 1;
}
.iq-tagline { color: #334155; font-size: .88em; margin-top: 6px; letter-spacing: .5px; }

.iq-section-header {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 24px 0 10px;
}
.iq-section-icon {
    width: 38px; height: 38px;
    background: linear-gradient(135deg, #6366f1, #8b5cf6);
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.1em;
    flex-shrink: 0;
    box-shadow: 0 4px 14px rgba(99,102,241,.35);
}
.iq-section-title { font-size: 1.15em; font-weight: 800; color: #f1f5f9; letter-spacing: -.3px; }
.iq-section-sub   { font-size: .78em; color: #475569; margin-top: 2px; }

.iq-card {
    background: linear-gradient(145deg, #0f172a, #111827);
    border: 1px solid #1e293b;
    border-radius: 16px;
    padding: 22px 26px;
    margin: 10px 0;
    transition: border-color .2s;
}
.iq-card:hover { border-color: #334155; }

.iq-info    { background:#0c1230;border-left:3px solid #6366f1;padding:12px 16px;border-radius:0 10px 10px 0;color:#a5b4fc;font-size:.84em;line-height:1.9;margin:8px 0; }
.iq-success { background:#061a0f;border-left:3px solid #22c55e;padding:12px 16px;border-radius:0 10px 10px 0;color:#86efac;font-size:.84em;line-height:1.9;margin:8px 0; }
.iq-warn    { background:#1a140a;border-left:3px solid #f59e0b;padding:12px 16px;border-radius:0 10px 10px 0;color:#fcd34d;font-size:.84em;line-height:1.9;margin:8px 0; }
.iq-danger  { background:#1a0808;border-left:3px solid #ef4444;padding:12px 16px;border-radius:0 10px 10px 0;color:#fca5a5;font-size:.84em;line-height:1.9;margin:8px 0; }

.iq-badge           { display:inline-block;padding:3px 12px;border-radius:20px;font-size:.72em;font-weight:700;letter-spacing:.5px; }
.iq-badge-purple    { background:#2e1065;color:#c4b5fd; }
.iq-badge-green     { background:#052e16;color:#86efac; }
.iq-badge-yellow    { background:#431407;color:#fcd34d; }
.iq-badge-red       { background:#450a0a;color:#fca5a5; }
.iq-badge-blue      { background:#0c1a4a;color:#93c5fd; }

.iq-divider { border:none;border-top:1px solid #1e293b;margin:20px 0; }

.iq-insight         { background:#0c1230;border-radius:10px;padding:14px 18px;margin:6px 0;border:1px solid #1e293b; }
.iq-insight-title   { color:#a5b4fc;font-size:.78em;font-weight:700;text-transform:uppercase;letter-spacing:.8px; }
.iq-insight-text    { color:#cbd5e1;font-size:.87em;margin-top:6px;line-height:1.7; }

.iq-col-header {
    color: #f1f5f9;
    font-size: 1em;
    font-weight: 700;
    padding: 10px 0 6px;
    border-bottom: 1px solid #1e293b;
    margin-bottom: 10px;
}

/* Feature card on landing page */
.iq-feature-card {
    background: linear-gradient(145deg, #0f172a, #111827);
    border: 1px solid #1e293b;
    border-top: 2px solid #6366f1;
    border-radius: 16px;
    padding: 28px 22px;
    text-align: center;
    height: 190px;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    transition: border-color .25s, transform .25s;
}
.iq-feature-card:hover {
    border-color: #8b5cf6;
    transform: translateY(-3px);
}
.iq-feature-icon  { font-size: 2em; margin-bottom: 10px; }
.iq-feature-title { color: #f1f5f9; font-weight: 800; font-size: .9em; margin-bottom: 8px; }
.iq-feature-desc  { color: #475569; font-size: .75em; line-height: 1.65; }

/* Supported file type card */
.iq-ftype-card {
    background: linear-gradient(145deg, #0f172a, #111827);
    border: 1px solid #1e293b;
    border-radius: 14px;
    padding: 14px 18px;
    margin: 6px 0;
}
.iq-ftype-cat   { color: #a5b4fc; font-size: .78em; font-weight: 700; margin-bottom: 6px; }
.iq-ftype-fmts  { color: #64748b; font-size: .76em; line-height: 1.8; }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# UTILITY HELPERS
# ════════════════════════════════════════════════════════════════

def ibox(t):  st.markdown(f'<div class="iq-info">💡 {t}</div>',    unsafe_allow_html=True)
def sbox(t):  st.markdown(f'<div class="iq-success">✅ {t}</div>', unsafe_allow_html=True)
def wbox(t):  st.markdown(f'<div class="iq-warn">⚠️ {t}</div>',   unsafe_allow_html=True)
def dbox(t):  st.markdown(f'<div class="iq-danger">🔴 {t}</div>', unsafe_allow_html=True)
def divider(): st.markdown('<hr class="iq-divider">', unsafe_allow_html=True)

def section(icon, title, sub=""):
    sub_html = f'<div class="iq-section-sub">{sub}</div>' if sub else ''
    st.markdown(f"""
    <div class="iq-section-header">
        <div class="iq-section-icon">{icon}</div>
        <div>
            <div class="iq-section-title">{title}</div>
            {sub_html}
        </div>
    </div>""", unsafe_allow_html=True)


def safe_float(v):
    try:
        f = float(v)
        return 0.0 if (np.isnan(f) or np.isinf(f)) else f
    except Exception:
        return 0.0


def safe_skew(series):
    try:
        n = pd.to_numeric(series, errors='coerce').dropna()
        return safe_float(n.skew()) if len(n) >= 3 else 0.0
    except Exception:
        return 0.0


def safe_numeric_stats(series):
    try:
        n = pd.to_numeric(series, errors='coerce').dropna()
        if len(n) == 0:
            return None
        q1, q3 = float(n.quantile(.25)), float(n.quantile(.75))
        iqr = q3 - q1
        outs = int(((n < q1 - 1.5 * iqr) | (n > q3 + 1.5 * iqr)).sum())
        return {
            'mean': safe_float(n.mean()), 'median': safe_float(n.median()),
            'std':  safe_float(n.std()),  'min':    safe_float(n.min()),
            'max':  safe_float(n.max()),  'skew':   safe_float(n.skew()),
            'q1': q1, 'q3': q3, 'iqr': iqr,
            'outliers': outs, 'count': len(n),
            'cv': safe_float(n.std() / (abs(n.mean()) + 1e-9) * 100),
        }
    except Exception:
        return None


def get_recommendation(col, df):
    try:
        present = df[col].dropna()
        pct     = df[col].isnull().sum() / max(len(df), 1) * 100
        dtype   = str(df[col].dtype)
        if len(present) == 0:
            return "Drop column — entirely null", "danger"
        if pct < 5:
            return f"Drop rows — only {pct:.1f}% missing, safe removal", "success"
        if dtype == 'object' or 'string' in dtype:
            mv = present.mode()[0] if len(present.mode()) > 0 else 'N/A'
            return f'Fill with Mode → "{mv}" (text column)', "info"
        num = pd.to_numeric(present, errors='coerce').dropna()
        if len(num) == 0:
            mv = present.mode()[0] if len(present.mode()) > 0 else 'N/A'
            return f'Fill with Mode → "{mv}" (non-numeric)', "info"
        skw = safe_skew(num)
        if abs(skw) > 1:
            return f"Fill with Median → {safe_float(num.median()):.4f} (skewed={skw:.2f})", "warn"
        return f"Fill with Mean → {safe_float(num.mean()):.4f} (symmetric={skw:.2f})", "info"
    except Exception:
        return "Use Mode or Custom fill", "info"


def compute_quality(df):
    try:
        if df is None or len(df) == 0 or df.shape[1] == 0:
            return 0, "🔴 No Data", 0.0, 0.0
        tc  = df.shape[0] * df.shape[1]
        mp  = safe_float(df.isnull().sum().sum() / max(tc, 1) * 100)
        dp  = safe_float(df.duplicated().sum() / max(len(df), 1) * 100)
        sc  = max(0, round(100 - min(mp * 2, 40) - min(dp * 2, 30)))
        gr  = ('🟢 Excellent' if sc >= 90 else '🟡 Good' if sc >= 70
               else '🟠 Fair' if sc >= 50 else '🔴 Needs Work')
        return sc, gr, mp, dp
    except Exception:
        return 0, "⚠️ Error", 0.0, 0.0


def fmt_bytes(b):
    for u in ['B', 'KB', 'MB', 'GB']:
        if b < 1024:
            return f"{b:.1f} {u}"
        b /= 1024
    return f"{b:.1f} TB"


# ════════════════════════════════════════════════════════════════
# UNIVERSAL FILE READER
# ════════════════════════════════════════════════════════════════

def read_any_file(file):
    name = file.name.lower()
    raw  = file.read()
    file.seek(0)
    meta = {
        'filename': file.name,
        'filesize': fmt_bytes(len(raw)),
        'filetype': 'Unknown',
        'notes': [],
        'ext': os.path.splitext(name)[1],
    }

    try:
        # ── CSV ──────────────────────────────────────────────────
        if name.endswith('.csv'):
            meta['filetype'] = 'CSV'
            for enc in ['utf-8', 'latin-1', 'cp1252', 'utf-16']:
                try:
                    df = pd.read_csv(io.BytesIO(raw), encoding=enc)
                    if enc != 'utf-8':
                        meta['notes'].append(f'Encoding: {enc}')
                    return df, meta, None
                except Exception:
                    continue
            return None, meta, 'CSV_ENCODE_ERROR'

        # ── EXCEL XLSX ───────────────────────────────────────────
        elif name.endswith('.xlsx'):
            meta['filetype'] = 'Excel (.xlsx)'
            xl = pd.ExcelFile(io.BytesIO(raw), engine='openpyxl')
            meta['sheets'] = xl.sheet_names
            if len(xl.sheet_names) == 1:
                return xl.parse(xl.sheet_names[0]), meta, None
            meta['excel_obj'] = xl
            return None, meta, 'MULTI_SHEET'

        # ── EXCEL XLS ────────────────────────────────────────────
        elif name.endswith('.xls'):
            meta['filetype'] = 'Excel (.xls)'
            xl = pd.ExcelFile(io.BytesIO(raw), engine='xlrd')
            meta['sheets'] = xl.sheet_names
            if len(xl.sheet_names) == 1:
                return xl.parse(xl.sheet_names[0]), meta, None
            meta['excel_obj'] = xl
            return None, meta, 'MULTI_SHEET'

        # ── JSON ─────────────────────────────────────────────────
        elif name.endswith('.json'):
            meta['filetype'] = 'JSON'
            try:
                df = pd.read_json(io.BytesIO(raw))
            except ValueError:
                df = pd.read_json(io.BytesIO(raw), lines=True)
                meta['notes'].append('Parsed as JSON Lines')
            return df, meta, None

        # ── PARQUET ──────────────────────────────────────────────
        elif name.endswith('.parquet'):
            meta['filetype'] = 'Parquet'
            return pd.read_parquet(io.BytesIO(raw)), meta, None

        # ── TSV / TXT ────────────────────────────────────────────
        elif name.endswith('.tsv') or name.endswith('.txt'):
            meta['filetype'] = 'TSV / Text'
            try:
                df = pd.read_csv(io.BytesIO(raw), sep='\t')
                meta['notes'].append('Tab-separated')
            except Exception:
                df = pd.read_csv(io.BytesIO(raw), sep=None, engine='python')
                meta['notes'].append('Auto-detected separator')
            return df, meta, None

        # ── SQLITE / DB ──────────────────────────────────────────
        elif name.endswith(('.db', '.sqlite', '.sqlite3')):
            meta['filetype'] = 'SQLite Database'
            tmp = f"_tmp_{file.name}"
            with open(tmp, 'wb') as f:
                f.write(raw)
            try:
                conn   = sqlite3.connect(tmp)
                tables = pd.read_sql("SELECT name FROM sqlite_master WHERE type='table'", conn)['name'].tolist()
                meta['tables']  = tables
                meta['db_path'] = tmp
                meta['db_conn'] = conn
                if not tables:
                    return None, meta, 'NO_TABLES'
                return None, meta, 'MULTI_TABLE'
            except Exception as e:
                try:
                    os.remove(tmp)
                except Exception:
                    pass
                return None, meta, str(e)

        # ── SQL SCRIPT ───────────────────────────────────────────
        elif name.endswith('.sql'):
            meta['filetype'] = 'SQL Script'
            for enc in ['utf-8', 'latin-1']:
                try:
                    meta['code'] = raw.decode(enc)
                    break
                except Exception:
                    continue
            return None, meta, 'CODE_FILE'

        # ── PYTHON ───────────────────────────────────────────────
        elif name.endswith('.py'):
            meta['filetype'] = 'Python Script'
            for enc in ['utf-8', 'latin-1']:
                try:
                    meta['code'] = raw.decode(enc)
                    break
                except Exception:
                    continue
            return None, meta, 'CODE_FILE'

        # ── PDF ──────────────────────────────────────────────────
        elif name.endswith('.pdf'):
            meta['filetype'] = 'PDF Document'
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(raw))
                pages  = []
                for i, pg in enumerate(reader.pages):
                    try:
                        pages.append(f"\n── Page {i+1} ──\n{pg.extract_text() or ''}")
                    except Exception:
                        pages.append(f"\n── Page {i+1} ── [unreadable]\n")
                meta['text']    = '\n'.join(pages)
                meta['n_pages'] = len(reader.pages)
                return None, meta, 'PDF_FILE'
            except ImportError:
                return None, meta, 'MISSING_LIB:PyPDF2'

        # ── WORD DOCX ────────────────────────────────────────────
        elif name.endswith('.docx'):
            meta['filetype'] = 'Word Document'
            try:
                from docx import Document
                doc    = Document(io.BytesIO(raw))
                paras  = [p.text for p in doc.paragraphs if p.text.strip()]
                tables = []
                for t in doc.tables:
                    rows = [[c.text.strip() for c in r.cells] for r in t.rows]
                    if len(rows) > 1:
                        try:
                            tables.append(pd.DataFrame(rows[1:], columns=rows[0]))
                        except Exception:
                            pass
                meta['paragraphs'] = paras
                meta['doc_tables'] = tables
                if tables:
                    meta['notes'].append(f'{len(tables)} table(s) extracted')
                    return tables[0], meta, None
                return None, meta, 'DOCX_TEXT'
            except ImportError:
                return None, meta, 'MISSING_LIB:python-docx'

        # ── IMAGE ────────────────────────────────────────────────
        elif name.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.webp', '.tiff')):
            meta['filetype'] = 'Image'
            try:
                from PIL import Image as PILImage
                img = PILImage.open(io.BytesIO(raw))
                meta['img']    = img
                meta['size']   = img.size
                meta['mode']   = img.mode
                meta['format'] = img.format
                return None, meta, 'IMAGE_FILE'
            except ImportError:
                return None, meta, 'MISSING_LIB:pillow'

        # ── POWER BI PBIX ────────────────────────────────────────
        elif name.endswith('.pbix'):
            meta['filetype'] = 'Power BI Report'
            try:
                with zipfile.ZipFile(io.BytesIO(raw)) as z:
                    contents = z.namelist()
                    meta['contents'] = contents
                    if 'Report/Layout' in contents:
                        try:
                            layout = json.loads(z.read('Report/Layout').decode('utf-8', 'replace'))
                            meta['pages'] = [s.get('displayName', '') for s in layout.get('sections', [])]
                        except Exception:
                            meta['pages'] = []
                    if 'Connections' in contents:
                        try:
                            meta['connections'] = z.read('Connections').decode('utf-8', 'replace')
                        except Exception:
                            pass
                    meta['data_files'] = [f for f in contents if 'DataModel' in f or 'data' in f.lower()]
                return None, meta, 'PBIX_FILE'
            except zipfile.BadZipFile:
                return None, meta, 'PBIX_CORRUPT'

        # ── XML ──────────────────────────────────────────────────
        elif name.endswith('.xml'):
            meta['filetype'] = 'XML'
            try:
                return pd.read_xml(io.BytesIO(raw)), meta, None
            except Exception as e:
                return None, meta, f'XML_ERROR:{e}'

        # ── ODS ──────────────────────────────────────────────────
        elif name.endswith('.ods'):
            meta['filetype'] = 'ODS Spreadsheet'
            return pd.read_excel(io.BytesIO(raw), engine='odf'), meta, None

        # ── FEATHER ──────────────────────────────────────────────
        elif name.endswith('.feather'):
            meta['filetype'] = 'Feather'
            return pd.read_feather(io.BytesIO(raw)), meta, None

        else:
            meta['filetype'] = 'Unsupported'
            return None, meta, 'UNKNOWN_TYPE'

    except Exception as e:
        meta['trace'] = traceback.format_exc()
        return None, meta, f'ERROR:{e}'


# ════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div style='text-align:center;padding:28px 0 16px;'>
        <div style='font-size:2.2em;font-weight:900;
            background:linear-gradient(135deg,#6366f1,#8b5cf6,#ec4899);
            -webkit-background-clip:text;-webkit-text-fill-color:transparent;
            letter-spacing:-1px;'>DataIQ</div>
        <div style='color:#334155;font-size:.72em;margin-top:4px;letter-spacing:1.5px;
            text-transform:uppercase;font-weight:600;'>Pro Analytics</div>
    </div>
    <hr class="iq-divider">
    <div style='color:#475569;font-size:.68em;font-weight:700;letter-spacing:1.5px;
        text-transform:uppercase;margin-bottom:12px;padding:0 4px;'>Navigation</div>
    """, unsafe_allow_html=True)

    steps = [
        ("📂", "Upload",   "File ingestion & parsing"),
        ("🔍", "Explore",  "Structure & statistics"),
        ("🔢", "Values",   "Frequency analysis"),
        ("📦", "GroupBy",  "Aggregation engine"),
        ("🧹", "Clean",    "Health check & repair"),
        ("📊", "Insights", "Intelligence & report"),
    ]
    for icon, label, desc in steps:
        st.markdown(f"""
        <div style='display:flex;align-items:center;gap:10px;padding:9px 12px;
            border-radius:10px;margin:3px 0;background:#0f172a;border:1px solid #1e293b;'>
            <div style='width:28px;height:28px;background:linear-gradient(135deg,#6366f1,#8b5cf6);
                border-radius:7px;display:flex;align-items:center;justify-content:center;
                font-size:.85em;flex-shrink:0;'>{icon}</div>
            <div>
                <div style='color:#e2e8f0;font-size:.82em;font-weight:700;'>{label}</div>
                <div style='color:#334155;font-size:.68em;'>{desc}</div>
            </div>
        </div>""", unsafe_allow_html=True)

    st.markdown('<hr class="iq-divider">', unsafe_allow_html=True)

    if st.session_state.get('loaded'):
        if st.button("🔄 New Dataset", key='sb_reset'):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()

        m = st.session_state.get('meta', {})
        if m:
            st.markdown(f"""
            <div style='background:#0f172a;border:1px solid #1e293b;border-radius:10px;padding:14px;'>
                <div style='color:#475569;font-size:.68em;font-weight:700;letter-spacing:1px;
                    text-transform:uppercase;margin-bottom:10px;'>Active File</div>
                <div style='color:#e2e8f0;font-size:.8em;font-weight:600;
                    word-break:break-all;'>{m.get('filename','')}</div>
                <div style='color:#475569;font-size:.72em;margin-top:4px;'>
                    {m.get('filetype','')} · {m.get('filesize','')}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("<div style='color:#1e293b;font-size:.68em;text-align:center;margin-top:20px;'>DataIQ Pro v3.0 · Enterprise Edition</div>", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# PAGE HEADER
# ════════════════════════════════════════════════════════════════
st.markdown("""
<div class="iq-page-header">
    <div class="iq-logo">DataIQ Pro</div>
    <div class="iq-tagline">
        Enterprise Data Intelligence Platform &nbsp;·&nbsp;
        Upload any file &nbsp;·&nbsp; Explore &nbsp;·&nbsp; Clean &nbsp;·&nbsp; Analyse &nbsp;·&nbsp; Report
    </div>
</div>""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# SECTION 1 — FILE UPLOAD
# ════════════════════════════════════════════════════════════════
section("📂", "Upload Dataset", "Supports 15+ file formats from any source")

# Supported types shown in a clean grid — outside the file uploader widget
SUPPORTED = {
    '📊 Tabular Data':   'CSV, Excel (.xlsx/.xls), JSON, Parquet, TSV, ODS, Feather, XML',
    '🗄️ Database':       'SQLite (.db, .sqlite, .sqlite3)',
    '💻 Code / Scripts': 'SQL (.sql), Python (.py)',
    '📄 Documents':      'PDF, Word (.docx)',
    '🖼️ Images':         'PNG, JPG, JPEG, BMP, GIF, WEBP, TIFF',
    '📈 BI Reports':     'Power BI (.pbix)',
}

with st.expander("📋 Supported File Types", expanded=False):
    cols = st.columns(3)
    for i, (cat, fmts) in enumerate(SUPPORTED.items()):
        with cols[i % 3]:
            st.markdown(f"""
            <div class="iq-ftype-card">
                <div class="iq-ftype-cat">{cat}</div>
                <div class="iq-ftype-fmts">{fmts}</div>
            </div>""", unsafe_allow_html=True)

st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

# ── File uploader ─────────────────────────────────────────────
file = st.file_uploader(
    label="Drop any file here — CSV, Excel, JSON, Parquet, SQLite, PDF, Word, Image, Power BI and more",
    type=['csv','xlsx','xls','json','parquet','tsv','txt','db','sqlite','sqlite3',
          'sql','py','pdf','docx','png','jpg','jpeg','bmp','gif','webp','tiff',
          'pbix','xml','ods','feather'],
    label_visibility='visible',
    key='file_uploader',
)

# ── LANDING PAGE ──────────────────────────────────────────────
if file is None:
    st.markdown("<br>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    features = [
        ("🔍", "Deep Exploration",  "Statistics, distributions, data types and full structural analysis"),
        ("🧹", "Smart Cleaning",    "Column-level diagnosis with inline one-click fixes"),
        ("🔗", "Intelligence",      "Correlations, outliers, cardinality and auto insights"),
        ("📋", "Executive Report",  "Auto-generated summary ready for stakeholders"),
    ]
    for col, (icon, title, desc) in zip([c1, c2, c3, c4], features):
        col.markdown(f"""
        <div class="iq-feature-card">
            <div class="iq-feature-icon">{icon}</div>
            <div class="iq-feature-title">{title}</div>
            <div class="iq-feature-desc">{desc}</div>
        </div>""", unsafe_allow_html=True)
    st.stop()


# ════════════════════════════════════════════════════════════════
# READ THE FILE (cache by filename to avoid re-parsing)
# ════════════════════════════════════════════════════════════════
if st.session_state.get('file_name') != file.name:
    with st.spinner("Parsing file…"):
        df_raw, meta, err = read_any_file(file)
    st.session_state.file_name  = file.name
    st.session_state.meta       = meta
    st.session_state.err        = err
    st.session_state.df_raw     = df_raw
    st.session_state.loaded     = True
    st.session_state.pop('cleaned_df', None)

meta   = st.session_state.meta
err    = st.session_state.err
df_raw = st.session_state.df_raw


# ════════════════════════════════════════════════════════════════
# NON-TABULAR FILE HANDLERS
# ════════════════════════════════════════════════════════════════

if err == 'MULTI_SHEET':
    sbox(f"<b>{meta['filename']}</b> — {len(meta['sheets'])} sheets detected")
    chosen = st.selectbox("Select sheet to analyse", meta['sheets'], key='sheet_pick')
    if st.button("Load Selected Sheet", key='load_sheet'):
        df_raw = meta['excel_obj'].parse(chosen)
        st.session_state.df_raw = df_raw
        st.session_state.err    = None
        meta['notes'].append(f"Sheet loaded: {chosen}")
        st.rerun()
    st.stop()

elif err == 'MULTI_TABLE':
    sbox(f"<b>{meta['filename']}</b> — SQLite database with {len(meta['tables'])} table(s)")
    chosen_tbl = st.selectbox("Select table to analyse", meta['tables'], key='tbl_pick')
    if st.button("Load Table", key='load_tbl'):
        conn   = sqlite3.connect(meta['db_path'])
        df_raw = pd.read_sql(f"SELECT * FROM [{chosen_tbl}]", conn)
        conn.close()
        try:
            os.remove(meta['db_path'])
        except Exception:
            pass
        st.session_state.df_raw = df_raw
        st.session_state.err    = None
        st.rerun()
    st.stop()

elif err == 'CODE_FILE':
    section("💻", f"{meta['filetype']} Viewer", meta['filename'])
    sbox(f"<b>{meta['filename']}</b> — {len(meta.get('code','').splitlines())} lines")
    lang = 'sql' if meta['filename'].lower().endswith('.sql') else 'python'
    st.code(meta.get('code', ''), language=lang)
    st.download_button("⬇️ Download File", data=meta.get('code', '').encode(),
                       file_name=meta['filename'], key='dl_code')
    st.stop()

elif err == 'PDF_FILE':
    section("📄", "PDF Document", meta['filename'])
    sbox(f"<b>{meta['filename']}</b> — {meta.get('n_pages', 0)} pages extracted")
    st.text_area("Extracted Text", meta.get('text', ''), height=500)
    st.download_button("⬇️ Download Text", data=meta.get('text', '').encode(),
                       file_name=meta['filename'].replace('.pdf', '.txt'), key='dl_pdf')
    st.stop()

elif err == 'DOCX_TEXT':
    section("📝", "Word Document", meta['filename'])
    sbox(f"<b>{meta['filename']}</b> — {len(meta.get('paragraphs', []))} paragraphs")
    for p in meta.get('paragraphs', []):
        st.markdown(f"<p style='color:#cbd5e1;font-size:.88em;line-height:1.8;padding:2px 0;'>{p}</p>",
                    unsafe_allow_html=True)
    st.stop()

elif err == 'IMAGE_FILE':
    section("🖼️", "Image File", meta['filename'])
    c1, c2 = st.columns([2, 1])
    with c1:
        st.image(meta['img'], use_container_width=True)
    with c2:
        st.markdown('<div class="iq-card">', unsafe_allow_html=True)
        for k, v in [("Filename", meta['filename']), ("Size", meta['filesize']),
                     ("Dimensions", f"{meta['size'][0]}×{meta['size'][1]} px"),
                     ("Mode", meta['mode']), ("Format", meta.get('format', 'N/A'))]:
            st.markdown(f"""
            <div style='display:flex;justify-content:space-between;
                padding:8px 0;border-bottom:1px solid #1e293b;'>
                <span style='color:#475569;font-size:.82em;'>{k}</span>
                <span style='color:#f1f5f9;font-size:.82em;font-weight:600;'>{v}</span>
            </div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

elif err == 'PBIX_FILE':
    section("📈", "Power BI Report", meta['filename'])
    sbox(f"<b>{meta['filename']}</b> — Power BI archive inspected")
    ibox("Power BI .pbix files store data in a proprietary compressed format. We can inspect the archive structure and metadata.")
    t1, t2, t3 = st.tabs(["📋 Report Pages", "🗂️ Archive Contents", "🔌 Connections"])
    with t1:
        pages = meta.get('pages', [])
        if pages:
            for i, p in enumerate(pages, 1):
                st.markdown(f"<div style='background:#0f172a;border:1px solid #1e293b;border-radius:8px;"
                            f"padding:12px 16px;margin:6px 0;color:#e2e8f0;font-size:.88em;'>"
                            f"<b style='color:#6366f1;'>Page {i}</b> &nbsp; {p}</div>",
                            unsafe_allow_html=True)
        else:
            wbox("Could not extract page names.")
    with t2:
        st.dataframe(pd.DataFrame({'File Path': meta.get('contents', [])}), use_container_width=True)
    with t3:
        st.code(meta.get('connections', 'No connection info found.')[:3000], language='json')
    st.stop()

elif err and err.startswith('MISSING_LIB:'):
    lib = err.split(':')[1]
    dbox(f"Missing library: <b>{lib}</b>. Install it with: <code>pip install {lib}</code>")
    st.stop()

elif err and err not in (None, 'MULTI_SHEET', 'MULTI_TABLE'):
    dbox(f"Could not process file: <b>{err}</b><br>"
         "Supported formats: CSV, Excel, JSON, Parquet, TSV, SQLite, SQL, Python, "
         "PDF, Word, Image, Power BI, XML, ODS, Feather")
    if meta.get('trace'):
        with st.expander("Technical Details"):
            st.code(meta['trace'])
    st.stop()

# ── Validate we have a DataFrame ──────────────────────────────
if df_raw is None or (hasattr(df_raw, 'empty') and df_raw.empty):
    dbox("File was read but produced no tabular data.")
    st.stop()

# ── Clean up column names ──────────────────────────────────────
df_raw.columns = df_raw.columns.astype(str).str.strip()

# ── Remove 100% empty columns ─────────────────────────────────
full_null = [c for c in df_raw.columns if df_raw[c].isnull().all()]
if full_null:
    df_raw = df_raw.drop(columns=full_null)
    wbox(f"Auto-removed {len(full_null)} fully-empty column(s): <b>{full_null}</b>")

if 'cleaned_df' not in st.session_state:
    st.session_state.cleaned_df = df_raw.copy()

data = df_raw.copy()   # immutable raw reference
df   = st.session_state.cleaned_df


# ════════════════════════════════════════════════════════════════
# FILE METADATA BANNER
# ════════════════════════════════════════════════════════════════
st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
notes_str = " &nbsp;·&nbsp; ".join(meta.get('notes', [])) if meta.get('notes') else ""
miss_clr  = "#ef4444" if data.isnull().sum().sum() > 0 else "#22c55e"
dupe_clr  = "#ef4444" if data.duplicated().sum() > 0   else "#22c55e"

st.markdown(f"""
<div style='background:linear-gradient(135deg,#0f172a,#0c1230);
    border:1px solid #1e293b;border-radius:14px;padding:18px 24px;
    display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:12px;'>
    <div style='display:flex;align-items:center;gap:10px;flex-wrap:wrap;'>
        <span style='color:#f1f5f9;font-weight:800;font-size:1em;'>{meta['filename']}</span>
        <span style='background:#1e1b4b;color:#a5b4fc;padding:3px 10px;
            border-radius:20px;font-size:.72em;font-weight:700;'>{meta['filetype']}</span>
        <span style='background:#1e293b;color:#64748b;padding:3px 10px;
            border-radius:20px;font-size:.72em;'>{meta['filesize']}</span>
    </div>
    {f"<div style='color:#475569;font-size:.76em;'>{notes_str}</div>" if notes_str else ""}
    <div style='display:flex;gap:18px;flex-wrap:wrap;'>
        <span style='color:#64748b;font-size:.78em;'><b style='color:#6366f1;'>{data.shape[0]:,}</b> rows</span>
        <span style='color:#64748b;font-size:.78em;'><b style='color:#6366f1;'>{data.shape[1]}</b> cols</span>
        <span style='color:#64748b;font-size:.78em;'><b style='color:{miss_clr};'>{data.isnull().sum().sum():,}</b> missing</span>
        <span style='color:#64748b;font-size:.78em;'><b style='color:{dupe_clr};'>{data.duplicated().sum():,}</b> dupes</span>
    </div>
</div>""", unsafe_allow_html=True)
st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# SECTION 2 — EXPLORE
# ════════════════════════════════════════════════════════════════
divider()
section("🔍", "Explore Dataset", "Structural analysis, statistics, distributions and quality overview")

k = st.columns(6)
k[0].metric("Total Rows",      f"{data.shape[0]:,}")
k[1].metric("Total Columns",   f"{data.shape[1]}")
k[2].metric("Numeric Cols",    f"{len(data.select_dtypes(include='number').columns)}")
k[3].metric("Text Cols",       f"{len(data.select_dtypes(include='object').columns)}")
k[4].metric("Missing Cells",   f"{data.isnull().sum().sum():,}")
k[5].metric("Duplicate Rows",  f"{data.duplicated().sum():,}")

st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

t1, t2, t3, t4, t5 = st.tabs(["📋 Raw Data", "📊 Statistics", "🔢 Data Types", "❓ Missing", "🔁 Duplicates"])

with t1:
    n = st.slider("Rows to preview", 5, min(2000, len(data)), min(25, len(data)), key='exp_n')
    st.dataframe(data.head(n), use_container_width=True)

with t2:
    nd = data.select_dtypes(include='number')
    cd = data.select_dtypes(include='object')
    if not nd.empty:
        st.markdown('<div class="iq-col-header">Numeric Columns</div>', unsafe_allow_html=True)
        try:
            st.dataframe(nd.describe().T.style.background_gradient(cmap='Blues').format(precision=3),
                         use_container_width=True)
        except Exception:
            st.dataframe(nd.describe().T, use_container_width=True)
    if not cd.empty:
        st.markdown('<div class="iq-col-header">Categorical Columns</div>', unsafe_allow_html=True)
        st.dataframe(cd.describe().T, use_container_width=True)
    if nd.empty and cd.empty:
        wbox("No columns available for statistics.")

with t3:
    dtype_rows = []
    for col in data.columns:
        samp = data[col].dropna()
        dtype_rows.append({
            'Column':         col,
            'Dtype':          str(data[col].dtype),
            'Non-Null':       int(data[col].notnull().sum()),
            'Null Count':     int(data[col].isnull().sum()),
            'Null %':         f"{data[col].isnull().sum()/max(len(data),1)*100:.1f}%",
            'Unique Values':  int(data[col].nunique()),
            'Sample':         str(samp.iloc[0]) if len(samp) > 0 else 'N/A',
        })
    st.dataframe(pd.DataFrame(dtype_rows), use_container_width=True)

with t4:
    miss_tab = pd.DataFrame({
        'Column':    data.columns,
        'Missing':   data.isnull().sum().values,
        'Present':   data.notnull().sum().values,
        'Missing %': (data.isnull().sum().values / max(len(data), 1) * 100).round(2),
    }).query('Missing > 0').sort_values('Missing %', ascending=False)
    if miss_tab.empty:
        sbox("Zero missing values — dataset is complete.")
    else:
        try:
            st.dataframe(
                miss_tab.style
                    .background_gradient(subset=['Missing %'], cmap='Reds')
                    .format({'Missing %': '{:.2f}%'}),
                use_container_width=True,
            )
        except Exception:
            st.dataframe(miss_tab, use_container_width=True)

with t5:
    dup_tab = data[data.duplicated(keep=False)]
    if dup_tab.empty:
        sbox("Zero duplicate rows — dataset is unique.")
    else:
        wbox(f"<b>{data.duplicated().sum():,}</b> duplicate rows detected "
             f"({data.duplicated().sum()/max(len(data),1)*100:.1f}%)")
        st.dataframe(dup_tab.sort_values(by=data.columns.tolist()).head(200),
                     use_container_width=True)


# ════════════════════════════════════════════════════════════════
# SECTION 3 — VALUE COUNTS
# ════════════════════════════════════════════════════════════════
divider()
section("🔢", "Value Count Analysis", "Complete frequency distribution for any column — no row limit")

with st.expander("Open Value Count Explorer"):
    vc_col = st.selectbox("Select column", options=list(data.columns), key='vc_col')
    if st.button("📊 Analyse Column", key='vc_btn'):
        try:
            vc = data[vc_col].value_counts(dropna=False).reset_index()
            vc.columns = ['Value', 'Count']
            vc['Value']        = vc['Value'].astype(str)
            vc['Percentage']   = (vc['Count'] / max(len(data), 1) * 100).round(4)
            vc['Cumulative %'] = vc['Percentage'].cumsum().round(2)

            v1, v2, v3, v4, v5 = st.columns(5)
            v1.metric("Total Rows",    f"{len(data):,}")
            v2.metric("Unique Values", f"{data[vc_col].nunique():,}")
            v3.metric("Missing",       f"{data[vc_col].isnull().sum():,}")
            v4.metric("Top Value",     str(vc.iloc[0]['Value'])[:20] if len(vc) > 0 else 'N/A')
            v5.metric("Top Count",     f"{vc.iloc[0]['Count']:,}" if len(vc) > 0 else '0')

            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            try:
                st.dataframe(
                    vc.style
                        .background_gradient(subset=['Count'], cmap='Blues')
                        .format({'Percentage': '{:.2f}%', 'Cumulative %': '{:.2f}%'}),
                    use_container_width=True,
                    height=min(600, 60 + 36 * len(vc)),
                )
            except Exception:
                st.dataframe(vc, use_container_width=True)

            sbox(f"Showing all <b>{len(vc):,}</b> unique values in <b>{vc_col}</b>")
        except Exception as e:
            dbox(f"Value count error: {e}")


# ════════════════════════════════════════════════════════════════
# SECTION 4 — GROUPBY
# ════════════════════════════════════════════════════════════════
divider()
section("📦", "GroupBy Analysis", "Aggregate and summarise your data across dimensions")

with st.expander("Open GroupBy Builder"):
    g1, g2, g3, g4 = st.columns(4)
    with g1:
        grp_cols = st.multiselect("Group By Columns", options=list(data.columns), key='grp_c')
    with g2:
        agg_col = st.selectbox("Aggregate Column", options=list(data.columns), key='agg_c')
    with g3:
        agg_op  = st.selectbox("Operation",
                      ['sum','mean','median','max','min','count','std','var','nunique'], key='agg_op')
    with g4:
        sort_res = st.selectbox("Sort By", ['Result (Desc)', 'Result (Asc)', 'Group'], key='sort_r')

    if grp_cols:
        try:
            gb = data.groupby(grp_cols, observed=True).agg(Result=(agg_col, agg_op)).reset_index()
            if sort_res == 'Result (Desc)':
                gb = gb.sort_values('Result', ascending=False)
            elif sort_res == 'Result (Asc)':
                gb = gb.sort_values('Result', ascending=True)

            r1, r2, r3, r4 = st.columns(4)
            r1.metric("Groups Found", f"{len(gb):,}")
            r2.metric("Operation",    agg_op.upper())
            is_num = pd.api.types.is_numeric_dtype(gb['Result'])
            r3.metric("Max Result",
                      f"{safe_float(gb['Result'].max()):.2f}" if is_num else str(gb['Result'].max())[:15])
            r4.metric("Min Result",
                      f"{safe_float(gb['Result'].min()):.2f}" if is_num else str(gb['Result'].min())[:15])

            st.dataframe(gb, use_container_width=True)
            sbox(f"<b>{agg_op.upper()}</b> of <b>{agg_col}</b> across "
                 f"<b>{', '.join(grp_cols)}</b> — {len(gb):,} groups")
        except Exception as e:
            dbox(f"GroupBy error: <b>{e}</b>")
    else:
        ibox("Select at least one column to group by.")


# ════════════════════════════════════════════════════════════════
# SECTION 5 — DATA HEALTH CHECK + INLINE CLEANING
# ════════════════════════════════════════════════════════════════
divider()
section("🧹", "Data Health Check & Inline Cleaning",
        "Every issue diagnosed and fixed at the same point — column by column")

df = st.session_state.cleaned_df

if len(df) == 0 or df.shape[1] == 0:
    dbox("Working dataset is empty.")
    if st.button("🔄 Reset to Original", key='empty_rst'):
        st.session_state.cleaned_df = data.copy()
        st.rerun()
    st.stop()

# ── Live health dashboard ──────────────────────────────────────
qs, gr, mpc, dpc = compute_quality(df)
hc = st.columns(7)
hc[0].metric("Rows",           f"{df.shape[0]:,}",  delta=f"{df.shape[0]-data.shape[0]:,}")
hc[1].metric("Columns",        f"{df.shape[1]}",    delta=f"{df.shape[1]-data.shape[1]}")
hc[2].metric("Missing Cells",  f"{df.isnull().sum().sum():,}")
hc[3].metric("Duplicate Rows", f"{df.duplicated().sum():,}")
hc[4].metric("Quality Score",  f"{qs}/100")
hc[5].metric("Grade",          gr)
hc[6].metric("Memory",         fmt_bytes(df.memory_usage(deep=True).sum()))

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

bar_w     = max(0, min(100, qs))
bar_color = "#22c55e" if qs >= 90 else "#f59e0b" if qs >= 70 else "#ef4444"
st.markdown(f"""
<div style='margin:4px 0 16px;'>
    <div style='display:flex;justify-content:space-between;margin-bottom:4px;'>
        <span style='color:#475569;font-size:.72em;font-weight:600;text-transform:uppercase;
            letter-spacing:.8px;'>Data Quality</span>
        <span style='color:{bar_color};font-size:.72em;font-weight:700;'>{qs}/100 — {gr}</span>
    </div>
    <div style='background:#1e293b;border-radius:6px;height:6px;'>
        <div style='width:{bar_w}%;height:6px;border-radius:6px;
            background:{bar_color};transition:width .5s;'></div>
    </div>
</div>""", unsafe_allow_html=True)

# ── 5A: DUPLICATE ROWS ────────────────────────────────────────
st.markdown('<div class="iq-card">', unsafe_allow_html=True)
st.markdown('<div class="iq-col-header">🔁 Duplicate Row Diagnosis</div>', unsafe_allow_html=True)

df = st.session_state.cleaned_df
dup_extra = df.duplicated(keep='first').sum()
dup_all   = df[df.duplicated(keep=False)]

if dup_extra == 0:
    sbox("No duplicate rows — dataset is fully unique.")
else:
    dup_pct = dup_extra / max(len(df), 1) * 100
    da, db, dc = st.columns(3)
    da.metric("Extra Copies to Remove", f"{dup_extra:,}")
    db.metric("Total Rows Involved",    f"{len(dup_all):,}")
    dc.metric("Duplication Rate",       f"{dup_pct:.2f}%")

    if   dup_pct < 1:  sbox(f"LOW — {dup_pct:.2f}%. Minimal impact.")
    elif dup_pct < 5:  wbox(f"MEDIUM — {dup_pct:.2f}%. Remove before analysis.")
    elif dup_pct < 20: wbox(f"HIGH — {dup_pct:.2f}%. Must clean before ML.")
    else:              dbox(f"CRITICAL — {dup_pct:.2f}%. Data pipeline issue suspected.")

    with st.expander("🔬 View Duplicated Rows"):
        st.dataframe(dup_all.sort_values(by=df.columns.tolist()).head(200), use_container_width=True)
        try:
            freq = dup_all.groupby(df.columns.tolist(), observed=True).size().reset_index(name='Copies')
            st.markdown("**Frequency of each duplicate group:**")
            st.dataframe(freq.sort_values('Copies', ascending=False), use_container_width=True)
        except Exception:
            pass

    ibox(f"<b>Recommendation: Remove {dup_extra:,} duplicate rows</b><br>"
         f"After removal: {len(df)-dup_extra:,} rows remain.<br>"
         "⚠️ Verify context: repeated transactions may be intentional — review before removing.")

    if st.button("🗑️ Remove All Duplicates", key='rm_dupe'):
        before = len(df)
        st.session_state.cleaned_df = df.drop_duplicates()
        sbox(f"Removed {before-len(st.session_state.cleaned_df):,} duplicate rows.")
        st.rerun()

st.markdown('</div>', unsafe_allow_html=True)
st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

# ── 5B: MISSING VALUES ────────────────────────────────────────
st.markdown('<div class="iq-card">', unsafe_allow_html=True)
st.markdown('<div class="iq-col-header">❓ Missing Value Diagnosis & Inline Fix</div>', unsafe_allow_html=True)

df = st.session_state.cleaned_df
mc = df.isnull().sum()
missing_cols = mc[mc > 0].index.tolist()

if not missing_cols:
    sbox("No missing values — dataset is complete.")
else:
    tm = mc.sum()
    tc = df.shape[0] * df.shape[1]
    ma, mb, mc_ = st.columns(3)
    ma.metric("Columns With Missing", f"{len(missing_cols)} / {df.shape[1]}")
    mb.metric("Total Missing Cells",  f"{tm:,}")
    mc_.metric("Overall Missing %",   f"{tm/max(tc,1)*100:.2f}%")

    diag = []
    for col in missing_cols:
        cnt = int(df[col].isnull().sum())
        pct = cnt / max(len(df), 1) * 100
        rec, _ = get_recommendation(col, df)
        diag.append({
            'Column':         col,
            'Type':           str(df[col].dtype),
            'Missing':        cnt,
            'Missing %':      f"{pct:.2f}%",
            'Severity':       ('🟢 Low' if pct < 5 else '🟡 Medium' if pct < 20
                                else '🟠 High' if pct < 50 else '🔴 Critical'),
            'Recommendation': rec,
        })
    st.dataframe(pd.DataFrame(diag), use_container_width=True)

    divider()
    st.markdown('<div class="iq-col-header">Fix Each Column Inline</div>', unsafe_allow_html=True)

    for col in missing_cols:
        cnt   = int(df[col].isnull().sum())
        pct   = cnt / max(len(df), 1) * 100
        dtype = str(df[col].dtype)
        pres  = df[col].dropna()
        rec, _ = get_recommendation(col, df)

        with st.expander(f"🔎  {col}  ·  {cnt:,} missing  ·  {pct:.1f}%  ·  {dtype}"):
            if   pct < 5:  sbox(f"LOW severity — {pct:.1f}% missing.")
            elif pct < 20: wbox(f"MEDIUM — {pct:.1f}% missing.")
            elif pct < 50: wbox(f"HIGH — {pct:.1f}% missing.")
            else:          dbox(f"CRITICAL — {pct:.1f}% missing. Consider dropping column.")

            ibox(f"<b>Recommended Strategy:</b> {rec}")

            stats = safe_numeric_stats(pres) if dtype != 'object' else None
            if stats:
                sc = st.columns(6)
                sc[0].metric("Mean",     f"{stats['mean']:.3f}")
                sc[1].metric("Median",   f"{stats['median']:.3f}")
                sc[2].metric("Std Dev",  f"{stats['std']:.3f}")
                sc[3].metric("Min",      f"{stats['min']:.3f}")
                sc[4].metric("Max",      f"{stats['max']:.3f}")
                sc[5].metric("Outliers", f"{stats['outliers']}")

            with st.expander(f"📄 View rows missing '{col}'"):
                null_rows = df[df[col].isnull()]
                st.dataframe(null_rows.head(100), use_container_width=True)
                if len(null_rows) > 100:
                    st.caption(f"Showing 100 of {len(null_rows)} null rows")

            divider()
            fc1, fc2 = st.columns([2, 1])
            with fc1:
                strat = st.selectbox(
                    "Choose fix strategy",
                    ['Drop rows with nulls', 'Fill with Mean', 'Fill with Median',
                     'Fill with Mode', 'Fill with Custom Value', 'Drop this column'],
                    key=f"s_{col}",
                )
            cval = ""
            with fc2:
                if strat == 'Fill with Custom Value':
                    cval = st.text_input("Custom value", placeholder="0 / Unknown / N/A", key=f"cv_{col}")

            if st.button(f"✅ Apply Fix · {col}", key=f"fix_{col}"):
                work = st.session_state.cleaned_df.copy()
                b    = work[col].isnull().sum()
                try:
                    if strat == 'Drop rows with nulls':
                        work = work.dropna(subset=[col])
                    elif strat == 'Fill with Mean':
                        mv = pd.to_numeric(work[col], errors='coerce').mean()
                        if pd.isna(mv):
                            wbox("No numeric values — cannot compute mean.")
                        else:
                            work[col] = work[col].fillna(mv)
                    elif strat == 'Fill with Median':
                        mv = pd.to_numeric(work[col], errors='coerce').median()
                        if pd.isna(mv):
                            wbox("No numeric values — cannot compute median.")
                        else:
                            work[col] = work[col].fillna(mv)
                    elif strat == 'Fill with Mode':
                        modes = work[col].mode()
                        if len(modes) == 0:
                            wbox("Cannot compute mode — column entirely null.")
                        else:
                            work[col] = work[col].fillna(modes[0])
                    elif strat == 'Fill with Custom Value':
                        if str(cval).strip() == '':
                            wbox("Enter a custom value first.")
                        else:
                            work[col] = work[col].fillna(cval)
                    elif strat == 'Drop this column':
                        work = work.drop(columns=[col])

                    st.session_state.cleaned_df = work
                    a = work[col].isnull().sum() if col in work.columns else 0
                    sbox(f"Fixed <b>{b-a:,}</b> cells in <b>{col}</b>. {a} remain.")
                    st.rerun()
                except Exception as e:
                    dbox(f"Fix failed: {e}")

    # Bulk fix
    divider()
    st.markdown('<div class="iq-col-header">⚡ Bulk Fix — All Missing At Once</div>', unsafe_allow_html=True)
    bfa, bfb = st.columns([3, 1])
    with bfa:
        bulk = st.selectbox("Bulk strategy", [
            'Numeric → Mean + Text → Mode',
            'Numeric → Median + Text → Mode',
            'Drop all rows with any null',
            'Fill everything with 0',
            'Fill everything with "Unknown"',
        ], key='bulk_strat')
    with bfb:
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
        if st.button("⚡ Apply Bulk Fix", key='bulk_fix'):
            work = st.session_state.cleaned_df.copy()
            try:
                if bulk == 'Drop all rows with any null':
                    work = work.dropna()
                elif bulk == 'Fill everything with 0':
                    work = work.fillna(0)
                elif bulk == 'Fill everything with "Unknown"':
                    work = work.fillna('Unknown')
                else:
                    for c in work.columns:
                        if work[c].isnull().sum() == 0:
                            continue
                        if str(work[c].dtype) == 'object':
                            modes = work[c].mode()
                            if len(modes) > 0:
                                work[c] = work[c].fillna(modes[0])
                        else:
                            num = pd.to_numeric(work[c], errors='coerce')
                            mv  = num.mean() if 'Mean' in bulk else num.median()
                            if not pd.isna(mv):
                                work[c] = work[c].fillna(mv)
                st.session_state.cleaned_df = work
                sbox("Bulk fix applied.")
                st.rerun()
            except Exception as e:
                dbox(f"Bulk fix error: {e}")

st.markdown('</div>', unsafe_allow_html=True)
st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

# ── 5C: DROP COLUMNS ──────────────────────────────────────────
st.markdown('<div class="iq-card">', unsafe_allow_html=True)
st.markdown('<div class="iq-col-header">❌ Drop Unwanted Columns</div>', unsafe_allow_html=True)

df = st.session_state.cleaned_df
suggest = []
for c in df.columns:
    u = df[c].nunique()
    n = df[c].isnull().sum() / max(len(df), 1) * 100
    if u == len(df):  suggest.append(f"{c}  [all unique — likely ID]")
    elif u == 1:      suggest.append(f"{c}  [only 1 unique value — zero variance]")
    elif n > 60:      suggest.append(f"{c}  [{n:.0f}% missing — mostly empty]")

if suggest:
    wbox("Columns worth considering for removal:<br>" +
         "<br>".join([f"&nbsp;&nbsp;&nbsp;• {s}" for s in suggest]))

dc1, dc2 = st.columns([3, 1])
with dc1:
    to_drop = st.multiselect("Select columns to remove", df.columns.tolist(), key='to_drop')
with dc2:
    st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
    if st.button("❌ Drop Selected", key='do_drop'):
        if to_drop:
            st.session_state.cleaned_df = df.drop(columns=to_drop)
            sbox(f"Dropped: {to_drop}")
            st.rerun()
        else:
            wbox("Select at least one column.")

st.markdown('</div>', unsafe_allow_html=True)
st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

# ── 5D: RESET + FINAL CLEANED DATASET ────────────────────────
st.markdown('<div class="iq-card">', unsafe_allow_html=True)
st.markdown('<div class="iq-col-header">✅ Final Cleaned Dataset</div>', unsafe_allow_html=True)

# Reset button
if st.button("🔄 Reset to Original Data", key='rst_clean'):
    st.session_state.cleaned_df = data.copy()
    st.rerun()

df = st.session_state.cleaned_df
qs2, gr2, _, _ = compute_quality(df)

fs = st.columns(6)
fs[0].metric("Final Rows",    f"{df.shape[0]:,}", delta=f"{df.shape[0]-data.shape[0]:,}")
fs[1].metric("Final Columns", f"{df.shape[1]}",   delta=f"{df.shape[1]-data.shape[1]}")
fs[2].metric("Missing",       f"{df.isnull().sum().sum():,}")
fs[3].metric("Duplicates",    f"{df.duplicated().sum():,}")
fs[4].metric("Quality",       f"{qs2}/100")
fs[5].metric("Grade",         gr2)

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
st.dataframe(df, use_container_width=True)

dl1, dl2, dl3 = st.columns(3)
with dl1:
    st.download_button(
        "⬇️ Download CSV",
        data=df.to_csv(index=False).encode('utf-8'),
        file_name="cleaned_data.csv",
        mime="text/csv",
        key='dl_csv',
    )
with dl2:
    try:
        buf = io.BytesIO()
        df.to_excel(buf, index=False, engine='openpyxl')
        st.download_button(
            "⬇️ Download Excel",
            data=buf.getvalue(),
            file_name="cleaned_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key='dl_xlsx',
        )
    except Exception:
        pass
with dl3:
    st.download_button(
        "⬇️ Download JSON",
        data=df.to_json(orient='records', indent=2).encode('utf-8'),
        file_name="cleaned_data.json",
        mime="application/json",
        key='dl_json',
    )

st.markdown('</div>', unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════
# SECTION 6 — AUTOMATED INTELLIGENCE
# ════════════════════════════════════════════════════════════════
divider()
section("📊", "Automated Data Intelligence",
        "Deep analysis, correlations, outliers and executive summary — auto-generated")

adf      = st.session_state.cleaned_df
num_cols = adf.select_dtypes(include='number').columns.tolist()
cat_cols = adf.select_dtypes(include='object').columns.tolist()

if len(adf) == 0:
    dbox("Cleaned dataset is empty. Reset and re-upload.")
    st.stop()

ov = st.columns(6)
ov[0].metric("Rows",         f"{adf.shape[0]:,}")
ov[1].metric("Columns",      f"{adf.shape[1]}")
ov[2].metric("Numeric",      f"{len(num_cols)}")
ov[3].metric("Categorical",  f"{len(cat_cols)}")
ov[4].metric("Memory",       fmt_bytes(adf.memory_usage(deep=True).sum()))
ov[5].metric("Completeness", f"{(1-adf.isnull().sum().sum()/max(adf.shape[0]*adf.shape[1],1))*100:.1f}%")

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

qs3, gr3, mp3, dp3 = compute_quality(adf)
qc = st.columns(4)
qc[0].metric("Quality Score", f"{qs3}/100")
qc[1].metric("Grade",         gr3)
qc[2].metric("Missing %",     f"{mp3:.2f}%")
qc[3].metric("Duplicate %",   f"{dp3:.2f}%")
ibox(f"Missing penalty: <b>-{min(mp3*2,40):.1f}pts</b> &nbsp;·&nbsp; "
     f"Duplicate penalty: <b>-{min(dp3*2,30):.1f}pts</b> &nbsp;·&nbsp; "
     f"Final Score: <b>{qs3}/100</b>")

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

# ── Numeric column intelligence ───────────────────────────────
if num_cols:
    st.markdown('<div class="iq-col-header">🔢 Numeric Column Intelligence</div>', unsafe_allow_html=True)
    for col in num_cols:
        with st.expander(f"📊  {col}"):
            stats = safe_numeric_stats(adf[col])
            if not stats:
                wbox(f"No valid numeric data in {col}")
                continue
            sc = st.columns(6)
            sc[0].metric("Mean",     f"{stats['mean']:.3f}")
            sc[1].metric("Median",   f"{stats['median']:.3f}")
            sc[2].metric("Std Dev",  f"{stats['std']:.3f}")
            sc[3].metric("Min",      f"{stats['min']:.3f}")
            sc[4].metric("Max",      f"{stats['max']:.3f}")
            sc[5].metric("Outliers", f"{stats['outliers']}")

            diff = abs(stats['mean'] - stats['median']) / (abs(stats['median']) + 1e-9) * 100
            skw  = stats['skew']
            out_pct = stats['outliers'] / max(stats['count'], 1) * 100
            ins = [
                f"Mean vs Median gap: {diff:.1f}% — {'⚠️ outliers pulling average' if diff>10 else '✅ symmetric distribution'}",
                f"Skewness: {skw:.3f} — {'📈 right-skewed' if skw>1 else '📉 left-skewed' if skw<-1 else '✅ approximately normal'}",
                f"CV: {stats['cv']:.1f}% — {'⚠️ high spread' if stats['cv']>50 else '✅ consistent'}",
                f"Outliers: {stats['outliers']} ({out_pct:.1f}%) — {'⚠️ review before ML' if stats['outliers']>0 else '✅ clean'}",
                f"Range: {stats['min']:.3f} → {stats['max']:.3f} (spread {stats['max']-stats['min']:.3f})",
                f"IQR: {stats['iqr']:.3f}  (Q1={stats['q1']:.3f}, Q3={stats['q3']:.3f})",
            ]
            st.markdown('<div class="iq-insight"><div class="iq-insight-title">Column Intelligence</div>'
                        '<div class="iq-insight-text">', unsafe_allow_html=True)
            for i in ins:
                st.markdown(f"<div style='padding:3px 0;color:#cbd5e1;font-size:.85em;'>• {i}</div>",
                            unsafe_allow_html=True)
            st.markdown("</div></div>", unsafe_allow_html=True)

# ── Categorical column intelligence ───────────────────────────
if cat_cols:
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown('<div class="iq-col-header">🔤 Categorical Column Intelligence</div>', unsafe_allow_html=True)
    for col in cat_cols:
        with st.expander(f"🏷️  {col}"):
            cd = adf[col].dropna()
            if len(cd) == 0:
                wbox(f"{col} has no non-null values.")
                continue
            vc    = cd.value_counts()
            uniq  = cd.nunique()
            top_v = str(vc.index[0]); top_c = int(vc.iloc[0])
            bot_v = str(vc.index[-1]); bot_c = int(vc.iloc[-1])
            dom   = top_c / max(len(cd), 1) * 100

            cc = st.columns(5)
            cc[0].metric("Unique Values", f"{uniq:,}")
            cc[1].metric("Top Value",     top_v[:18])
            cc[2].metric("Top Count",     f"{top_c:,}")
            cc[3].metric("Dominance",     f"{dom:.1f}%")
            cc[4].metric("Missing",       f"{adf[col].isnull().sum():,}")

            vc_df = vc.head(20).reset_index()
            vc_df.columns = ['Value', 'Count']
            vc_df['%'] = (vc_df['Count'] / max(len(cd), 1) * 100).round(2)
            st.dataframe(vc_df, use_container_width=True)

            ins = []
            if uniq == 1:
                ins.append(f"⚠️ Only one value ('{top_v}') — zero variance. Drop this column.")
            elif uniq >= len(cd):
                ins.append("⚠️ All values unique — likely an ID/key column. Drop before ML.")
            elif uniq <= 5:
                ins.append(f"✅ Low cardinality ({uniq}) — ideal for one-hot encoding.")
            elif uniq > 50:
                ins.append(f"⚠️ High cardinality ({uniq}) — use target or frequency encoding.")
            else:
                ins.append(f"✅ Moderate cardinality ({uniq} values).")

            if dom > 80:
                ins.append(f"⚠️ '{top_v}' dominates at {dom:.1f}% — class imbalance risk in ML.")
            elif dom > 50:
                ins.append(f"ℹ️ '{top_v}' is majority at {dom:.1f}%.")
            else:
                ins.append(f"✅ Balanced distribution. Top value holds only {dom:.1f}%.")
            ins.append(f"Rarest: '{bot_v}' with {bot_c} occurrence(s).")

            st.markdown('<div class="iq-insight"><div class="iq-insight-title">Column Intelligence</div>'
                        '<div class="iq-insight-text">', unsafe_allow_html=True)
            for i in ins:
                st.markdown(f"<div style='padding:3px 0;color:#cbd5e1;font-size:.85em;'>• {i}</div>",
                            unsafe_allow_html=True)
            st.markdown("</div></div>", unsafe_allow_html=True)

# ── Correlation intelligence ───────────────────────────────────
pairs_df = pd.DataFrame()
if len(num_cols) >= 2:
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown('<div class="iq-col-header">🔗 Correlation Intelligence</div>', unsafe_allow_html=True)
    try:
        corr  = adf[num_cols].apply(pd.to_numeric, errors='coerce').corr()
        pairs = []
        for i in range(len(corr.columns)):
            for j in range(i + 1, len(corr.columns)):
                v = corr.iloc[i, j]
                if not (np.isnan(v) or np.isinf(v)):
                    av = abs(float(v))
                    pairs.append({
                        'Column A':    corr.columns[i],
                        'Column B':    corr.columns[j],
                        'Correlation': round(float(v), 4),
                        'Abs':         round(av, 4),
                        'Strength':    ('🔴 Strong' if av >= .7 else '🟡 Moderate' if av >= .4 else '⚪ Weak'),
                        'Direction':   ('↗ Positive' if v > 0 else '↘ Negative'),
                    })
        if pairs:
            pairs_df = pd.DataFrame(pairs).sort_values('Abs', ascending=False).drop(columns=['Abs'])
            st.dataframe(pairs_df, use_container_width=True)

            strong = pairs_df[pairs_df['Correlation'].abs() >= .7]
            if not strong.empty:
                st.markdown('<div class="iq-col-header" style="font-size:.85em;">Strong Correlations</div>',
                            unsafe_allow_html=True)
                for _, row in strong.iterrows():
                    ibox(f"<b>{row['Column A']}</b> ↔ <b>{row['Column B']}</b> = "
                         f"{row['Correlation']:.4f} &nbsp; {row['Strength']} &nbsp; {row['Direction']}")
            else:
                ibox("No strong correlations (|r| ≥ 0.7) found.")
    except Exception as e:
        wbox(f"Correlation error: {e}")

# ── Outlier intelligence ───────────────────────────────────────
out_rows = []
if num_cols:
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown('<div class="iq-col-header">🚨 Outlier Intelligence</div>', unsafe_allow_html=True)
    for col in num_cols:
        try:
            stats = safe_numeric_stats(adf[col])
            if not stats or stats['count'] < 4:
                continue
            op = stats['outliers'] / max(stats['count'], 1) * 100
            out_rows.append({
                'Column':        col,
                'Outlier Count': stats['outliers'],
                'Outlier %':     round(op, 2),
                'Lower Fence':   round(stats['q1'] - 1.5 * stats['iqr'], 3),
                'Upper Fence':   round(stats['q3'] + 1.5 * stats['iqr'], 3),
                'Status':        ('🔴 Investigate' if op > 5 else '🟡 Monitor' if op > 0 else '🟢 Clean'),
            })
        except Exception:
            continue
    if out_rows:
        st.dataframe(
            pd.DataFrame(out_rows).sort_values('Outlier %', ascending=False),
            use_container_width=True,
        )

# ── Executive summary report ───────────────────────────────────
st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
st.markdown('<div class="iq-col-header">📋 Executive Summary Report</div>', unsafe_allow_html=True)

ts    = datetime.now().strftime("%Y-%m-%d %H:%M")
lines = [
    "DataIQ Pro — Executive Summary Report",
    f"Generated : {ts}",
    f"Source    : {meta['filename']} ({meta['filetype']})",
    "",
    "── DATASET OVERVIEW ──────────────────────────────────────",
    f"Dimensions         : {adf.shape[0]:,} rows × {adf.shape[1]} columns",
    f"Numeric columns    : {len(num_cols)}",
    f"Categorical columns: {len(cat_cols)}",
    f"Memory usage       : {fmt_bytes(adf.memory_usage(deep=True).sum())}",
    "",
    "── DATA QUALITY ──────────────────────────────────────────",
    f"Quality Score  : {qs3}/100 — {gr3}",
    (f"Missing cells  : {adf.isnull().sum().sum():,} ({mp3:.2f}%)"
     if adf.isnull().sum().sum() > 0 else "Missing cells  : None"),
    (f"Duplicate rows : {adf.duplicated().sum():,}"
     if adf.duplicated().sum() > 0 else "Duplicate rows : None"),
    "",
    "── INSIGHTS ──────────────────────────────────────────────",
]

if not pairs_df.empty:
    strong_n = len(pairs_df[pairs_df['Correlation'].abs() >= .7])
    lines.append(f"Strong correlations (|r|≥0.7) : {strong_n}")

if out_rows:
    worst = max(out_rows, key=lambda x: x['Outlier %'])
    if worst['Outlier %'] > 5:
        lines.append(f"Highest outlier column : {worst['Column']} at {worst['Outlier %']}%")

if cat_cols:
    hc_list = [c for c in cat_cols if adf[c].nunique() > 50]
    if hc_list:
        lines.append(f"High-cardinality columns : {', '.join(hc_list)}")

single_val = [c for c in adf.columns if adf[c].nunique() == 1]
if single_val:
    lines.append(f"Zero-variance columns    : {', '.join(single_val)}")

id_cols = [c for c in adf.columns if adf[c].nunique() == len(adf)]
if id_cols:
    lines.append(f"Likely ID columns        : {', '.join(id_cols)}")

lines += ["", "── RECOMMENDATIONS ───────────────────────────────────────"]
if adf.isnull().sum().sum() > 0:
    lines.append("• Address remaining missing values before ML training")
if adf.duplicated().sum() > 0:
    lines.append("• Remove remaining duplicate rows")
if out_rows and any(r['Outlier %'] > 5 for r in out_rows):
    lines.append("• Investigate high-outlier columns — consider capping or removal")
if not pairs_df.empty and len(pairs_df[pairs_df['Correlation'].abs() >= .7]) > 0:
    lines.append("• Review strongly correlated columns — consider dimensionality reduction")
lines.append("• Validate high-cardinality columns before one-hot encoding")

report_text = "\n".join(lines)

st.markdown(f"""
<div style='background:#0c1a0c;border:1px solid #166534;border-radius:14px;
    padding:24px 28px;font-family:monospace;font-size:.82em;color:#86efac;
    line-height:1.9;white-space:pre-wrap;max-height:420px;overflow-y:auto;'>
{report_text}
</div>""", unsafe_allow_html=True)

dr1, dr2 = st.columns(2)
with dr1:
    st.download_button(
        "⬇️ Download Report (.txt)",
        data=report_text.encode('utf-8'),
        file_name=f"DataIQ_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
        mime="text/plain",
        key='dl_report',
    )
with dr2:
    st.download_button(
        "⬇️ Download Report (.md)",
        data=report_text.encode('utf-8'),
        file_name=f"DataIQ_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
        mime="text/markdown",
        key='dl_report_md',
    )


# ════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════
divider()
st.markdown(f"""
<div style='text-align:center;padding:16px 0 8px;'>
    <div style='font-size:1.1em;font-weight:900;
        background:linear-gradient(135deg,#6366f1,#8b5cf6,#ec4899);
        -webkit-background-clip:text;-webkit-text-fill-color:transparent;
        letter-spacing:-.5px;'>DataIQ Pro</div>
    <div style='color:#1e293b;font-size:.72em;margin-top:6px;letter-spacing:.5px;'>
        Enterprise Analytics Platform &nbsp;·&nbsp; v3.1 &nbsp;·&nbsp;
        Built with Streamlit &amp; Pandas &nbsp;·&nbsp;
        {datetime.now().strftime("%Y")}
    </div>
</div>""", unsafe_allow_html=True)