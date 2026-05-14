"""
DataIQ Pro Analytics - Universal File Handler

This module provides universal file reading capabilities for 20+ file formats.
"""

import io
import pandas as pd
import numpy as np
import streamlit as st
from typing import Tuple, Dict, Any, Optional
import os
import tempfile
import zipfile
import sqlite3
import json
import chardet


def read_any_file(file) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """
    Universal file reader that handles 20+ formats.
    Returns (dataframe_or_None, meta_dict, error_code)
    
    Args:
        file: The uploaded file object from Streamlit.
    
    Returns:
        Tuple of (DataFrame or None, metadata dict, error code string).
    """
    meta = {
        'filename': getattr(file, 'name', 'unknown'),
        'filesize': getattr(file, 'size', 0),
        'filetype': '',
        'ext': '',
        'notes': [],
        'trace': ''
    }
    
    try:
        # Get file extension
        filename = meta['filename'].lower()
        if '.' in filename:
            meta['ext'] = filename.split('.')[-1]
        else:
            meta['ext'] = 'noext'
        
        meta['filetype'] = meta['ext'].upper()
        
        # Read file content
        file.seek(0)
        content = file.read()
        file.seek(0)  # Reset for pandas reading
        
        # Handle different formats
        if meta['ext'] in ['csv', 'tsv', 'txt']:
            return _read_text_file(file, content, meta)
        elif meta['ext'] in ['xlsx', 'xls']:
            return _read_excel_file(file, meta)
        elif meta['ext'] == 'json':
            return _read_json_file(file, content, meta)
        elif meta['ext'] == 'parquet':
            return _read_parquet_file(file, meta)
        elif meta['ext'] in ['db', 'sqlite', 'sqlite3']:
            return _read_sqlite_file(file, meta)
        elif meta['ext'] == 'sql':
            return _read_sql_file(content, meta)
        elif meta['ext'] in ['py', 'python']:
            return _read_python_file(content, meta)
        elif meta['ext'] == 'pdf':
            return _read_pdf_file(content, meta)
        elif meta['ext'] == 'docx':
            return _read_docx_file(file, meta)
        elif meta['ext'] in ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'webp', 'tiff']:
            return _read_image_file(file, meta)
        elif meta['ext'] == 'pbix':
            return _read_pbix_file(file, meta)
        elif meta['ext'] == 'xml':
            return _read_xml_file(file, meta)
        elif meta['ext'] == 'ods':
            return _read_ods_file(file, meta)
        elif meta['ext'] == 'feather':
            return _read_feather_file(file, meta)
        else:
            meta['notes'].append(f"Unknown file extension: {meta['ext']}")
            return None, meta, "UNKNOWN_TYPE"
            
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Exception during reading: {str(e)}")
        return None, meta, "READ_ERROR"


def _read_text_file(file, content: bytes, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle CSV, TSV, TXT files with encoding detection."""
    encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    
    for encoding in encodings_to_try:
        try:
            text_content = content.decode(encoding)
            
            # Detect separator
            if meta['ext'] == 'tsv':
                sep = '\t'
            elif meta['ext'] == 'csv':
                # Try to detect separator
                sample = text_content[:1024]
                if '\t' in sample and sample.count('\t') > sample.count(','):
                    sep = '\t'
                else:
                    sep = ','
            else:  # txt
                # Auto-detect separator
                sample = text_content[:1024]
                separators = [',', '\t', ';', '|']
                sep_counts = [(s, sample.count(s)) for s in separators]
                sep = max(sep_counts, key=lambda x: x[1])[0]
            
            # Read with pandas using StringIO
            from io import StringIO
            df = pd.read_csv(StringIO(text_content), sep=sep, engine='python')
            
            # Remove fully empty columns
            df = df.dropna(axis=1, how='all')
            
            meta['notes'].append(f"Successfully read with {encoding} encoding, separator '{sep}'")
            return df, meta, ""
            
        except Exception as e:
            meta['notes'].append(f"Failed with {encoding}: {str(e)}")
            continue
    
    meta['notes'].append("All encodings failed")
    return None, meta, "CSV_ENCODE_ERROR"


def _read_excel_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Excel files (.xlsx, .xls)."""
    try:
        if meta['ext'] == 'xlsx':
            excel_file = pd.ExcelFile(file, engine='openpyxl')
        else:  # xls
            excel_file = pd.ExcelFile(file, engine='xlrd')
        
        sheet_names = excel_file.sheet_names
        
        if len(sheet_names) == 1:
            df = pd.read_excel(excel_file, sheet_name=sheet_names[0])
            meta['notes'].append(f"Read single sheet: {sheet_names[0]}")
            return df, meta, ""
        else:
            meta['excel_obj'] = excel_file
            meta['sheet_names'] = sheet_names
            meta['notes'].append(f"Multiple sheets found: {sheet_names}")
            return None, meta, "MULTI_SHEET"
            
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Excel reading failed: {str(e)}")
        return None, meta, "EXCEL_ERROR"


def _read_json_file(file, content: bytes, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle JSON files."""
    try:
        # Try normal JSON first
        text_content = content.decode('utf-8')
        data = json.loads(text_content)
        
        # Try to convert to DataFrame
        df = pd.json_normalize(data) if isinstance(data, list) else pd.DataFrame([data])
        
        meta['notes'].append("Read as standard JSON")
        return df, meta, ""
        
    except Exception:
        try:
            # Try JSON Lines format
            file.seek(0)
            df = pd.read_json(file, lines=True)
            meta['notes'].append("Read as JSON Lines format")
            return df, meta, ""
            
        except Exception as e:
            meta['trace'] = str(e)
            meta['notes'].append(f"JSON reading failed: {str(e)}")
            return None, meta, "JSON_ERROR"


def _read_parquet_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Parquet files."""
    try:
        df = pd.read_parquet(file)
        meta['notes'].append("Successfully read Parquet file")
        return df, meta, ""
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Parquet reading failed: {str(e)}")
        return None, meta, "PARQUET_ERROR"


def _read_sqlite_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle SQLite database files."""
    temp_path = None
    try:
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.db') as temp_file:
            temp_file.write(file.read())
            temp_path = temp_file.name
        
        # Connect and get table list
        conn = sqlite3.connect(temp_path)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = [row[0] for row in cursor.fetchall()]
        conn.close()
        
        if not tables:
            meta['notes'].append("No tables found in database")
            return None, meta, "NO_TABLES"
        
        meta['tables'] = tables
        meta['db_path'] = temp_path
        meta['notes'].append(f"Found tables: {tables}")
        return None, meta, "MULTI_TABLE"
        
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"SQLite reading failed: {str(e)}")
        return None, meta, "SQLITE_ERROR"
    finally:
        # Cleanup will be handled by caller if needed
        pass


def _read_sql_file(content: bytes, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle SQL files."""
    try:
        text_content = content.decode('utf-8')
        meta['code'] = text_content
        meta['notes'].append("SQL file detected - contains code/text")
        return None, meta, "CODE_FILE"
    except Exception as e:
        try:
            text_content = content.decode('latin-1')
            meta['code'] = text_content
            meta['notes'].append("SQL file decoded with latin-1")
            return None, meta, "CODE_FILE"
        except Exception as e2:
            meta['trace'] = str(e2)
            meta['notes'].append(f"SQL file reading failed: {str(e2)}")
            return None, meta, "SQL_ERROR"


def _read_python_file(content: bytes, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Python files."""
    try:
        text_content = content.decode('utf-8')
        meta['code'] = text_content
        meta['notes'].append("Python file detected - contains code")
        return None, meta, "CODE_FILE"
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Python file reading failed: {str(e)}")
        return None, meta, "PYTHON_ERROR"


def _read_pdf_file(content: bytes, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle PDF files."""
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        meta['text'] = text
        meta['n_pages'] = len(reader.pages)
        meta['notes'].append(f"Extracted text from {len(reader.pages)} pages")
        return None, meta, "PDF_FILE"
        
    except ImportError:
        meta['notes'].append("PyPDF2 not installed")
        return None, meta, "MISSING_LIB:PyPDF2"
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"PDF reading failed: {str(e)}")
        return None, meta, "PDF_ERROR"


def _read_docx_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Word documents."""
    try:
        from docx import Document
        
        doc = Document(file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Check for tables
        tables_data = []
        for table in doc.tables:
            table_df = pd.DataFrame()
            for i, row in enumerate(table.rows):
                row_data = [cell.text for cell in row.cells]
                if i == 0:
                    table_df = pd.DataFrame(columns=row_data)
                else:
                    table_df.loc[len(table_df)] = row_data
            if not table_df.empty:
                tables_data.append(table_df)
        
        if tables_data:
            df = tables_data[0]  # Return first table
            meta['notes'].append(f"Found {len(tables_data)} tables, returning first one")
            return df, meta, ""
        else:
            meta['paragraphs'] = paragraphs
            meta['notes'].append(f"No tables found, extracted {len(paragraphs)} paragraphs")
            return None, meta, "DOCX_TEXT"
            
    except ImportError:
        meta['notes'].append("python-docx not installed")
        return None, meta, "MISSING_LIB:python-docx"
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Word reading failed: {str(e)}")
        return None, meta, "DOCX_ERROR"


def _read_image_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle image files."""
    try:
        from PIL import Image
        import io
        
        image = Image.open(file)
        meta['pil_image'] = image
        meta['size'] = image.size
        meta['mode'] = image.mode
        meta['format'] = image.format
        meta['notes'].append(f"Image: {image.format} {image.size} {image.mode}")
        return None, meta, "IMAGE_FILE"
        
    except ImportError:
        meta['notes'].append("Pillow not installed")
        return None, meta, "MISSING_LIB:Pillow"
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Image reading failed: {str(e)}")
        return None, meta, "IMAGE_ERROR"


def _read_pbix_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Power BI files."""
    try:
        import zipfile
        import json
        
        pbix_zip = zipfile.ZipFile(file)
        contents = pbix_zip.namelist()
        
        # Try to extract Report/Layout
        report_data = None
        connections_data = None
        
        if 'Report/Layout' in contents:
            with pbix_zip.open('Report/Layout') as f:
                report_content = f.read().decode('utf-8')
                report_data = json.loads(report_content)
        
        if 'Connections' in contents:
            with pbix_zip.open('Connections') as f:
                connections_content = f.read().decode('utf-8')
                connections_data = json.loads(connections_content)
        
        meta['contents'] = contents
        meta['report_data'] = report_data
        meta['connections_data'] = connections_data
        
        if report_data and 'sections' in report_data:
            page_names = [section.get('displayName', 'Unknown') for section in report_data['sections']]
            meta['page_names'] = page_names
            meta['notes'].append(f"Power BI file with pages: {page_names}")
        else:
            meta['notes'].append("Power BI file contents extracted")
        
        return None, meta, "PBIX_FILE"
        
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Power BI reading failed: {str(e)}")
        return None, meta, "PBIX_ERROR"


def _read_xml_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle XML files."""
    try:
        df = pd.read_xml(file)
        meta['notes'].append("Successfully read XML file")
        return df, meta, ""
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"XML reading failed: {str(e)}")
        return None, meta, "XML_ERROR"


def _read_ods_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle ODS files."""
    try:
        df = pd.read_excel(file, engine='odf')
        meta['notes'].append("Successfully read ODS file")
        return df, meta, ""
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"ODS reading failed: {str(e)}")
        return None, meta, "ODS_ERROR"


def _read_feather_file(file, meta: Dict[str, Any]) -> Tuple[Optional[pd.DataFrame], Dict[str, Any], str]:
    """Handle Feather files."""
    try:
        df = pd.read_feather(file)
        meta['notes'].append("Successfully read Feather file")
        return df, meta, ""
    except Exception as e:
        meta['trace'] = str(e)
        meta['notes'].append(f"Feather reading failed: {str(e)}")
        return None, meta, "FEATHER_ERROR"
    